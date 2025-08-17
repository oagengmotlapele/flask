from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document
doc = Document()

# Set page margins to fit everything on one A4 page
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Add placeholder for logo
logo_paragraph = doc.add_paragraph("Company Logo Here")
logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = logo_paragraph.runs[0]
run.font.size = Pt(10)
run.font.italic = True

# Title
title = doc.add_heading('Quotation: Scholarship Scraping Website Development', 1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Company Info
doc.add_paragraph("Company: 65Bugs")
doc.add_paragraph("Currency: Botswana Pula (BWP)")
doc.add_paragraph("Date: July 23, 2025")

# Summary
doc.add_heading('Quotation Summary', level=2)
doc.add_paragraph(
    "This quotation covers the development of a dynamic scholarship scraping website. "
    "The system includes live data scraping, admin and user dashboards, and essential features for scholarship discovery."
)

# Quotation Table
doc.add_heading('Quotation Details', level=2)

quotation_items = [
    ("Dynamic Scraper (1 site)", "BWP 2,000"),
    ("Database Design & Setup", "BWP 1,000"),
    ("User Login & Authentication", "BWP 1,000"),
    ("Admin Dashboard", "BWP 1,200"),
    ("Frontend UI/UX", "BWP 800"),
    ("Search & Filtering Features", "BWP 800"),
    ("Scheduler Setup", "BWP 600"),
    ("Deployment & Security", "BWP 600"),
    ("Annual Hosting (.com/.co.bw domain)", "BWP 1,000"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.autofit = True

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Item Description'
hdr_cells[1].text = 'Cost'

# Data rows
for item, price in quotation_items:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = price

# Total row
total_row = table.add_row().cells
total_row[0].text = "Total Quotation"
total_row[1].text = "BWP 9,000"

# Maintenance
doc.add_heading('Maintenance & Hosting', level=2)
doc.add_paragraph(
    "• Annual Hosting & Domain Fee: BWP 1,000\n"
    "• Scraper Update Fee (for structure changes): BWP 200 per update\n"
    "*Note: Maintenance fees are billed separately when required.*"
)

# Payment Terms
doc.add_heading('Payment Terms', level=2)
doc.add_paragraph(
    "• 50% upfront: BWP 4,500\n"
    "• 50% upon completion: BWP 4,500\n"
    "*Annual hosting and maintenance charges not included in initial payment.*"
)

# Footer
footer = doc.add_paragraph("Thank you for choosing 65Bugs.")
footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save
output_path = "65Bugs_Quotation_A4_With_Logo_Place.docx"
doc.save(output_path)

output_path
